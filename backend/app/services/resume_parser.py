import re
import nltk
import spacy
from typing import Dict, List, Optional, Tuple, Any
import PyPDF2
from docx import Document
import io
import hashlib
from datetime import datetime
import phonenumbers
from nameparser import HumanName
import logging

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

try:
    nltk.data.find('taggers/averaged_perceptron_tagger')
except LookupError:
    nltk.download('averaged_perceptron_tagger')

try:
    nltk.data.find('chunkers/maxent_ne_chunker')
except LookupError:
    nltk.download('maxent_ne_chunker')

try:
    nltk.data.find('corpora/words')
except LookupError:
    nltk.download('words')

class AdvancedResumeParser:
    def __init__(self):
        """Initialize the comprehensive resume parser with enhanced NLP models and pattern recognition."""
        self.logger = logging.getLogger(__name__)
        
        # Load spaCy model (fallback if not available)
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            self.logger.warning("spaCy model not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Comprehensive section header patterns
        self.section_patterns = {
            'personal_info': [
                r'personal\s+(?:information|info|details)',
                r'contact\s+(?:information|info|details)',
                r'contact\s+me',
                r'about\s+me',
                r'profile'
            ],
            'summary': [
                r'(?:career\s+)?summary',
                r'(?:professional\s+)?summary',
                r'(?:career\s+)?objective',
                r'profile',
                r'overview',
                r'about\s+me',
                r'statement',
                r'personal\s+statement'
            ],
            'education': [
                r'education(?:al\s+background)?',
                r'academic\s+(?:background|qualifications|history)',
                r'qualifications',
                r'degrees?',
                r'academics'
            ],
            'experience': [
                r'(?:work\s+)?experience',
                r'professional\s+experience',
                r'employment\s+(?:history|background)',
                r'career\s+(?:history|background)',
                r'work\s+history',
                r'internship(?:s)?'
            ],
            'skills': [
                r'(?:technical\s+)?skills',
                r'core\s+competencies',
                r'expertise',
                r'proficiencies',
                r'key\s+skills',
                r'technologies',
                r'areas\s+of\s+expertise'
            ],
            'projects': [
                r'projects?',
                r'portfolio',
                r'selected\s+projects',
                r'research\s+work',
                r'thesis',
                r'dissertation',
                r'case\s+stud(?:y|ies)'
            ],
            'certifications': [
                r'certifications?',
                r'certificates?',
                r'professional\s+certifications?',
                r'credentials?',
                r'licenses?'
            ],
            'achievements': [
                r'achievements?',
                r'awards?',
                r'honors?',
                r'accomplishments?',
                r'recognition',
                r'hackathons?',
                r'competitions?'
            ],
            'courses': [
                r'courses?',
                r'training',
                r'workshops?',
                r'bootcamps?',
                r'programs?',
                r'simulations?'
            ],
            'publications': [
                r'publications?',
                r'research\s+papers?',
                r'journals?',
                r'conference\s+proceedings?'
            ],
            'volunteer': [
                r'volunteer(?:ing)?',
                r'community\s+service',
                r'social\s+work'
            ],
            'languages': [
                r'languages?',
                r'language\s+skills'
            ],
            'references': [
                r'references?',
                r'mentors?'
            ],
            'interests': [
                r'interests?',
                r'hobbies',
                r'personal\s+interests',
                r'activities'
            ]
        }
        
        # Comprehensive technical skills database with categories and synonyms
        self.tech_skills_database = {
            'programming_languages': {
                'keywords': [
                    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c', 'go', 'rust', 
                    'kotlin', 'swift', 'php', 'ruby', 'r', 'matlab', 'scala', 'perl', 'objective-c',
                    'dart', 'haskell', 'shell', 'bash', 'powershell', 'lua', 'assembly', 'fortran',
                    'cobol', 'ada', 'lisp', 'prolog', 'erlang', 'clojure', 'sql', 'nosql'
                ],
                'synonyms': {
                    'js': 'javascript',
                    'ts': 'typescript',
                    'py': 'python',
                    'c++': 'cpp',
                    'c#': 'csharp',
                    'objective-c': 'objc'
                }
            },
            'web_development': {
                'keywords': [
                    'html', 'css', 'sass', 'scss', 'less', 'bootstrap', 'tailwind', 'tailwindcss',
                    'jquery', 'react', 'angular', 'vue', 'vue.js', 'svelte', 'ember', 'backbone',
                    'node.js', 'express', 'express.js', 'next.js', 'nuxt.js', 'gatsby', 'webpack',
                    'vite', 'parcel', 'rollup', 'graphql', 'rest', 'grpc', 'openapi', 'swagger'
                ],
                'synonyms': {
                    'nodejs': 'node.js',
                    'expressjs': 'express.js',
                    'nextjs': 'next.js',
                    'nuxtjs': 'nuxt.js',
                    'vuejs': 'vue.js',
                    'reactjs': 'react'
                }
            },
            'frameworks_libraries': {
                'keywords': [
                    'django', 'flask', 'fastapi', 'spring', 'spring boot', 'laravel', 'rails',
                    'ruby on rails', 'asp.net', 'codeigniter', 'koa', 'tensorflow', 'pytorch',
                    'keras', 'scikit-learn', 'xgboost', 'lightgbm', 'catboost', 'hugging face',
                    'transformers', 'spacy', 'nltk', 'opencv', 'pandas', 'numpy', 'scipy',
                    'statsmodels', 'matplotlib', 'seaborn', 'plotly'
                ],
                'synonyms': {
                    'sklearn': 'scikit-learn',
                    'tf': 'tensorflow',
                    'hf': 'hugging face',
                    'cv2': 'opencv'
                }
            },
            'databases': {
                'keywords': [
                    'mysql', 'postgresql', 'sqlite', 'oracle', 'sql server', 'mariadb',
                    'mongodb', 'cassandra', 'couchdb', 'dynamodb', 'redis', 'firestore',
                    'snowflake', 'bigquery', 'redshift', 'neo4j', 'influxdb', 'fauna'
                ],
                'synonyms': {
                    'postgres': 'postgresql',
                    'mongo': 'mongodb',
                    'dynamo': 'dynamodb'
                }
            },
            'cloud_platforms': {
                'keywords': [
                    'aws', 'amazon web services', 'azure', 'microsoft azure', 'gcp', 'google cloud',
                    'google cloud platform', 'ibm cloud', 'digitalocean', 'heroku', 'vercel',
                    'netlify', 'ec2', 's3', 'lambda', 'rds', 'cloudformation', 'compute engine',
                    'firebase', 'pub/sub', 'ai platform', 'app service', 'cosmosdb', 'aks'
                ],
                'synonyms': {
                    'amazon web services': 'aws',
                    'google cloud platform': 'gcp',
                    'microsoft azure': 'azure'
                }
            },
            'devops_cicd': {
                'keywords': [
                    'docker', 'kubernetes', 'jenkins', 'gitlab ci/cd', 'circleci', 'travis ci',
                    'argocd', 'helm', 'ansible', 'puppet', 'chef', 'terraform', 'vagrant',
                    'prometheus', 'grafana', 'elk stack', 'elasticsearch', 'logstash', 'kibana',
                    'splunk', 'github actions'
                ],
                'synonyms': {
                    'k8s': 'kubernetes',
                    'tf': 'terraform',
                    'elk': 'elk stack'
                }
            },
            'version_control': {
                'keywords': [
                    'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial', 'jira', 'trello',
                    'asana', 'confluence', 'slack', 'microsoft teams', 'teams'
                ],
                'synonyms': {
                    'ms teams': 'microsoft teams'
                }
            },
            'testing_qa': {
                'keywords': [
                    'selenium', 'junit', 'pytest', 'mocha', 'chai', 'jasmine', 'postman',
                    'newman', 'cucumber', 'testng', 'appium', 'cypress', 'unit testing',
                    'integration testing', 'end-to-end testing', 'e2e testing'
                ],
                'synonyms': {
                    'e2e': 'end-to-end testing'
                }
            },
            'mobile_development': {
                'keywords': [
                    'ios', 'android', 'react native', 'flutter', 'xamarin', 'ionic', 'cordova',
                    'phonegap', 'swift ui', 'swiftui', 'jetpack compose'
                ],
                'synonyms': {
                    'rn': 'react native'
                }
            },
            'data_science_ai': {
                'keywords': [
                    'machine learning', 'deep learning', 'artificial intelligence', 'neural networks',
                    'data mining', 'data analysis', 'statistics', 'regression', 'classification',
                    'clustering', 'nlp', 'natural language processing', 'computer vision',
                    'reinforcement learning', 'data science', 'big data'
                ],
                'synonyms': {
                    'ml': 'machine learning',
                    'dl': 'deep learning',
                    'ai': 'artificial intelligence',
                    'cv': 'computer vision',
                    'rl': 'reinforcement learning'
                }
            },
            'cybersecurity': {
                'keywords': [
                    'cybersecurity', 'information security', 'infosec', 'kali linux', 'wireshark',
                    'metasploit', 'burp suite', 'nessus', 'owasp', 'iso 27001', 'gdpr', 'hipaa',
                    'soc 2', 'penetration testing', 'ethical hacking'
                ],
                'synonyms': {
                    'infosec': 'information security',
                    'pentest': 'penetration testing'
                }
            },
            'business_tools': {
                'keywords': [
                    'excel', 'power bi', 'tableau', 'qlikview', 'google analytics', 'sap',
                    'salesforce', 'zoho', 'ms office', 'microsoft office', 'google workspace',
                    'office 365'
                ],
                'synonyms': {
                    'powerbi': 'power bi',
                    'ms office': 'microsoft office',
                    'o365': 'office 365'
                }
            }
        }
        
        # Soft skills database
        self.soft_skills = [
            'communication', 'leadership', 'teamwork', 'collaboration', 'problem-solving',
            'analytical thinking', 'creativity', 'adaptability', 'critical thinking',
            'decision-making', 'time management', 'conflict resolution', 'public speaking',
            'negotiation', 'project management', 'strategic thinking', 'innovation',
            'emotional intelligence', 'mentoring', 'coaching'
        ]
        
        # Language proficiency levels
        self.proficiency_levels = [
            'native', 'fluent', 'advanced', 'intermediate', 'basic', 'beginner',
            'conversational', 'professional', 'limited', 'expert'
        ]
        
        # Common languages
        self.languages = [
            'english', 'spanish', 'french', 'german', 'italian', 'portuguese', 'chinese',
            'mandarin', 'japanese', 'korean', 'arabic', 'hindi', 'russian', 'dutch',
            'swedish', 'norwegian', 'polish', 'turkish', 'hebrew', 'thai', 'vietnamese'
        ]
    
    def calculate_file_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of file content for duplicate detection."""
        return hashlib.sha256(content).hexdigest()
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file with multiple methods for better extraction."""
        text = ""
        
        try:
            # Method 1: PyPDF2 for standard text extraction
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text + "\n"
                    
                    # Also try to extract text with different orientations
                    if page.rotation != 0:
                        try:
                            rotated_page = page.rotate(-page.rotation)
                            rotated_text = rotated_page.extract_text()
                            if len(rotated_text) > len(page_text):
                                text += f"\n--- Page {page_num + 1} (Rotated) ---\n"
                                text += rotated_text + "\n"
                        except:
                            pass
                            
                except Exception as e:
                    self.logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            # Method 2: Try alternative text extraction if minimal text found
            if len(text.strip()) < 100:
                try:
                    import pdfplumber
                    with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                except ImportError:
                    self.logger.info("pdfplumber not available, using PyPDF2 only")
                except Exception as e:
                    self.logger.warning(f"pdfplumber extraction failed: {e}")
            
            # Clean up the extracted text
            text = self._clean_extracted_text(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF text: {e}")
            # Try fallback method
            try:
                import fitz  # PyMuPDF
                pdf_document = fitz.open(stream=file_content, filetype="pdf")
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]
                    text += page.get_text() + "\n"
                pdf_document.close()
            except ImportError:
                self.logger.warning("PyMuPDF not available for fallback extraction")
            except Exception as fallback_error:
                self.logger.error(f"Fallback PDF extraction failed: {fallback_error}")
                
        return text

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file with comprehensive content extraction."""
        text = ""
        
        try:
            doc = Document(io.BytesIO(file_content))
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                text += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "--- End Table ---\n"
            
            # Extract text from headers and footers
            for section in doc.sections:
                # Headers
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text += f"[HEADER] {paragraph.text}\n"
                
                # Footers
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text += f"[FOOTER] {paragraph.text}\n"
            
            # Extract text from text boxes and shapes (if any)
            try:
                from docx.oxml.ns import qn
                from docx.oxml import parse_xml
                
                # This is a more advanced extraction for text boxes
                for element in doc.element.iter():
                    if element.tag.endswith('}txbxContent'):
                        for paragraph in element.iter():
                            if paragraph.tag.endswith('}t'):
                                if paragraph.text:
                                    text += f"[TEXTBOX] {paragraph.text}\n"
            except Exception as e:
                self.logger.debug(f"Advanced DOCX extraction failed: {e}")
            
            # Clean up the extracted text
            text = self._clean_extracted_text(text)
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {e}")
            
        return text

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = []
        for line in text.split('\n'):
            cleaned_line = ' '.join(line.split())  # Normalize whitespace
            if cleaned_line:  # Skip empty lines
                lines.append(cleaned_line)
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(lines)
        
        # Remove redundant page markers
        cleaned_text = re.sub(r'\n--- Page \d+ ---\n', '\n', cleaned_text)
        
        # Remove excessive newlines (more than 2 consecutive)
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        return cleaned_text.strip()
    
    def extract_personal_info(self, text: str) -> Dict[str, Any]:
        """Extract comprehensive personal information with structured format."""
        personal_info = {
            'name': {'first': '', 'middle': '', 'last': '', 'full': ''},
            'emails': [],
            'phones': [],
            'location': {'address': '', 'city': '', 'state': '', 'country': '', 'zip_code': ''},
            'links': {
                'linkedin': '',
                'github': '',
                'portfolio': '',
                'personal_website': '',
                'researchgate': '',
                'blog': '',
                'other_links': []
            }
        }
        
        # Extract name using multiple methods
        name_info = self._extract_name(text)
        personal_info['name'].update(name_info)
        
        # Extract emails
        emails = self._extract_emails(text)
        personal_info['emails'] = emails
        
        # Extract phone numbers
        phones = self._extract_phones(text)
        personal_info['phones'] = phones
        
        # Extract location information
        location_info = self._extract_location(text)
        personal_info['location'].update(location_info)
        
        # Extract links and categorize them
        links_info = self._extract_links(text)
        personal_info['links'].update(links_info)
        
        return personal_info
    
    def _extract_name(self, text: str) -> Dict[str, str]:
        """Extract full name with first, middle, last name breakdown."""
        name_info = {'first': '', 'middle': '', 'last': '', 'full': ''}
        
        # Method 1: spaCy NER for PERSON entities
        if self.nlp:
            doc = self.nlp(text[:1000])  # Process first 1000 chars
            for ent in doc.ents:
                if ent.label_ == "PERSON" and len(ent.text.split()) >= 2:
                    name_info['full'] = ent.text.strip()
                    break
        
        # Method 2: Look in first few lines for name patterns
        if not name_info['full']:
            lines = text.split('\n')[:10]
            for line in lines:
                line_clean = line.strip()
                # Skip lines with contact info, email, phone
                if any(indicator in line_clean.lower() for indicator in ['@', 'phone', 'email', 'linkedin', 'github']):
                    continue
                
                # Look for name patterns (2-4 words, mostly alphabetic)
                words = line_clean.split()
                if 2 <= len(words) <= 4 and all(word.replace('.', '').isalpha() for word in words):
                    name_info['full'] = line_clean
                    break
        
        # Parse name components using nameparser
        if name_info['full']:
            try:
                parsed_name = HumanName(name_info['full'])
                name_info['first'] = parsed_name.first
                name_info['middle'] = parsed_name.middle
                name_info['last'] = parsed_name.last
            except Exception as e:
                self.logger.debug(f"Name parsing error: {e}")
                # Fallback: split by spaces
                name_parts = name_info['full'].split()
                if len(name_parts) >= 2:
                    name_info['first'] = name_parts[0]
                    name_info['last'] = name_parts[-1]
                    if len(name_parts) > 2:
                        name_info['middle'] = ' '.join(name_parts[1:-1])
        
        return name_info
    
    def _extract_emails(self, text: str) -> List[str]:
        """Extract and normalize email addresses."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        # Normalize emails (lowercase, trim)
        normalized_emails = [email.lower().strip() for email in emails]
        unique_emails = list(dict.fromkeys(normalized_emails))  # Remove duplicates while preserving order
        
        return unique_emails
    
    def _extract_phones(self, text: str) -> List[str]:
        """Extract and normalize phone numbers with international support."""
        phones = []
        
        # Method 1: Use phonenumbers library for robust extraction
        try:
            for match in phonenumbers.PhoneNumberMatcher(text, "US"):
                formatted = phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.E164)
                phones.append(formatted)
        except Exception as e:
            self.logger.debug(f"Phonenumbers library error: {e}")
        
        # Method 2: Regex patterns for various phone formats
        phone_patterns = [
            r'\+?1?[-.\s]?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})',  # US format
            r'\+?([0-9]{1,4})[-.\s]?([0-9]{3,4})[-.\s]?([0-9]{3,4})[-.\s]?([0-9]{3,4})',  # International
            r'(\d{10})',  # 10 digits
            r'\+?(\d{11,15})'  # International with country code
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    phone = ''.join(match)
                else:
                    phone = match
                
                # Validate phone length
                if 10 <= len(re.sub(r'\D', '', phone)) <= 15:
                    phones.append(phone)
        
        # Remove duplicates and return
        unique_phones = list(dict.fromkeys(phones))
        return unique_phones
        unique_phones = list(dict.fromkeys(phones))
        
        phone_info = {
            'primary': unique_phones[0] if unique_phones else '',
            'secondary': unique_phones[1:] if len(unique_phones) > 1 else [],
            'all_phones': unique_phones
        }
        
        return phone_info
    
    def _extract_location(self, text: str) -> Dict[str, str]:
        """Extract location information with granular details."""
        location_info = {'address': '', 'city': '', 'state': '', 'country': '', 'zip_code': ''}
        
        # Location keywords
        location_keywords = ['address', 'location', 'city', 'state', 'country', 'zip', 'postal']
        
        lines = text.split('\n')
        for line in lines:
            line_lower = line.lower()
            
            # Check if line contains location keywords
            if any(keyword in line_lower for keyword in location_keywords):
                location_text = line.strip()
                
                # Extract ZIP/Postal code
                zip_pattern = r'\b(\d{5}(-\d{4})?|\d{6}|[A-Z]\d[A-Z]\s?\d[A-Z]\d)\b'
                zip_match = re.search(zip_pattern, location_text)
                if zip_match:
                    location_info['zip_code'] = zip_match.group(1)
                
                # Use spaCy for location entity recognition
                if self.nlp:
                    doc = self.nlp(location_text)
                    for ent in doc.ents:
                        if ent.label_ in ['GPE', 'LOC']:  # Geopolitical entity or location
                            if not location_info['city']:
                                location_info['city'] = ent.text
                            elif not location_info['state']:
                                location_info['state'] = ent.text
                            elif not location_info['country']:
                                location_info['country'] = ent.text
                
                # Store full address if not already captured
                if not location_info['address']:
                    location_info['address'] = location_text
                
                break
        
        return location_info
    
    def _extract_links(self, text: str) -> Dict[str, Any]:
        """Extract and categorize all links from the resume."""
        links_info = {
            'linkedin': '',
            'github': '',
            'portfolio': '',
            'personal_website': '',
            'researchgate': '',
            'blog': '',
            'other_links': []
        }
        
        # URL pattern
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+|www\.[^\s<>"{}|\\^`\[\]]+|[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}(?:/[^\s]*)?'
        urls = re.findall(url_pattern, text)
        
        # Also look for domain mentions without http
        domain_pattern = r'(?:linkedin\.com/in/|github\.com/|researchgate\.net/profile/|medium\.com/@)([A-Za-z0-9_.-]+)'
        domain_matches = re.findall(domain_pattern, text, re.IGNORECASE)
        
        for url in urls:
            url_lower = url.lower()
            
            if 'linkedin.com' in url_lower:
                links_info['linkedin'] = url if url.startswith('http') else f'https://{url}'
            elif 'github.com' in url_lower:
                links_info['github'] = url if url.startswith('http') else f'https://{url}'
            elif 'researchgate.net' in url_lower:
                links_info['researchgate'] = url if url.startswith('http') else f'https://{url}'
            elif any(blog_platform in url_lower for blog_platform in ['medium.com', 'blog.', 'wordpress.', 'blogspot.']):
                links_info['blog'] = url if url.startswith('http') else f'https://{url}'
            elif any(portfolio_indicator in url_lower for portfolio_indicator in ['portfolio', 'site', 'personal']):
                links_info['portfolio'] = url if url.startswith('http') else f'https://{url}'
            else:
                # Check if it's a personal website (not a major platform)
                major_platforms = ['google.com', 'facebook.com', 'twitter.com', 'instagram.com', 'youtube.com']
                if not any(platform in url_lower for platform in major_platforms):
                    if not links_info['personal_website']:
                        links_info['personal_website'] = url if url.startswith('http') else f'https://{url}'
                    else:
                        links_info['other_links'].append(url if url.startswith('http') else f'https://{url}')
        
        return links_info
    
    def extract_summary(self, text: str) -> str:
        """Extract career summary/objective with improved detection."""
        summary = ""
        
        lines = text.split('\n')
        in_summary_section = False
        summary_lines = []
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check if line is a summary header
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns['summary']):
                in_summary_section = True
                continue
            
            # If in summary section, collect lines until next section
            if in_summary_section:
                # Check if we hit another major section
                other_sections = ['education', 'experience', 'skills', 'projects', 'certifications']
                if any(any(re.search(pattern, line_lower) for pattern in self.section_patterns[section]) 
                       for section in other_sections):
                    break
                
                # Add meaningful content lines
                if len(line_clean.split()) > 3:  # Filter out very short lines
                    summary_lines.append(line_clean)
        
        # If explicit summary section found
        if summary_lines:
            summary = ' '.join(summary_lines)
        else:
            # Fallback: look for substantial paragraphs in the first part of resume
            for i, line in enumerate(lines[:15]):  # First 15 lines
                line_clean = line.strip()
                
                # Skip contact info lines
                if any(indicator in line_clean.lower() for indicator in 
                       ['@', 'phone', 'linkedin', 'github', 'email']):
                    continue
                
                # Look for substantial text (likely summary)
                if len(line_clean.split()) > 10 and '.' in line_clean:
                    summary = line_clean
                    # Check if next lines are continuation
                    for j in range(i + 1, min(len(lines), i + 5)):
                        next_line = lines[j].strip()
                        if (len(next_line.split()) > 5 and 
                            not any(section_word in next_line.lower() for section_word in 
                                   ['education', 'experience', 'skills', 'work'])):
                            summary += ' ' + next_line
                        else:
                            break
                    break
        
        return summary.strip()

    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract comprehensive education information with detailed parsing."""
        education_entries = []
        
        lines = text.split('\n')
        in_education_section = False
        current_entry = {}
        
        # Degree patterns
        degree_patterns = [
            r'(?:bachelor|master|doctorate|phd|mba|b\.?a\.?|b\.?s\.?|b\.?tech|b\.?e\.?|m\.?a\.?|m\.?s\.?|m\.?tech|ph\.?d\.?)',
            r'(?:associate|diploma|certificate|certification)',
            r'(?:undergraduate|graduate|postgraduate)',
            r'(?:high\s+school|secondary\s+school|12th|10th)'
        ]
        
        # Institution keywords
        institution_keywords = [
            'university', 'college', 'institute', 'school', 'academy', 'polytechnic',
            'technology', 'engineering', 'management', 'business', 'medical'
        ]
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check if entering education section
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns['education']):
                in_education_section = True
                continue
            
            # Check if leaving education section
            if in_education_section:
                other_sections = ['experience', 'skills', 'projects', 'certifications', 'achievements']
                if any(any(re.search(pattern, line_lower) for pattern in self.section_patterns[section]) 
                       for section in other_sections):
                    # Save current entry if exists
                    if current_entry:
                        education_entries.append(current_entry)
                        current_entry = {}
                    break
            
            if in_education_section:
                # Check for degree patterns
                degree_found = False
                for pattern in degree_patterns:
                    if re.search(pattern, line_lower):
                        # Save previous entry if exists
                        if current_entry:
                            education_entries.append(current_entry)
                        
                        # Start new entry
                        current_entry = {
                            'institution': '',
                            'degree': line_clean,
                            'field_of_study': '',
                            'start_year': '',
                            'end_year': '',
                            'duration': '',
                            'cgpa': '',
                            'percentage': '',
                            'grade': '',
                            'coursework': [],
                            'specialization': ''
                        }
                        degree_found = True
                        break
                
                # If degree found, look for additional info in nearby lines
                if degree_found:
                    # Extract dates from current line
                    self._extract_education_dates(line_clean, current_entry)
                    
                    # Look in surrounding lines for institution and other details
                    search_range = range(max(0, i-2), min(len(lines), i+4))
                    for j in search_range:
                        if j == i:
                            continue
                        
                        nearby_line = lines[j].strip()
                        nearby_lower = nearby_line.lower()
                        
                        # Look for institution
                        if (any(keyword in nearby_lower for keyword in institution_keywords) and 
                            not current_entry['institution']):
                            current_entry['institution'] = nearby_line
                        
                        # Look for CGPA/percentage
                        self._extract_education_scores(nearby_line, current_entry)
                        
                        # Look for coursework
                        if any(keyword in nearby_lower for keyword in ['coursework', 'subjects', 'major', 'minor']):
                            current_entry['coursework'].append(nearby_line)
                
                # Look for institution lines (without degree keywords)
                elif (any(keyword in line_lower for keyword in institution_keywords) and 
                      current_entry and not current_entry['institution']):
                    current_entry['institution'] = line_clean
                
                # Look for additional details for current entry
                elif current_entry:
                    # Check for scores
                    self._extract_education_scores(line_clean, current_entry)
                    
                    # Check for coursework
                    if any(keyword in line_lower for keyword in ['coursework', 'relevant courses', 'subjects']):
                        current_entry['coursework'].append(line_clean)
        
        # Save final entry
        if current_entry:
            education_entries.append(current_entry)
        
        # Clean up entries
        cleaned_entries = []
        for entry in education_entries:
            if entry.get('degree') or entry.get('institution'):
                # Clean coursework
                if entry['coursework']:
                    entry['coursework'] = [course.strip() for course in entry['coursework'] if course.strip()]
                cleaned_entries.append(entry)
        
        return cleaned_entries
    
    def _extract_education_dates(self, text: str, entry: Dict[str, Any]):
        """Extract dates from education text."""
        # Date patterns for education
        date_patterns = [
            r'(\d{4})\s*[-–—]\s*(\d{4})',  # 2018-2022
            r'(\d{4})\s*[-–—]\s*(present|current)',  # 2020-Present
            r'(\d{1,2}/\d{4})\s*[-–—]\s*(\d{1,2}/\d{4})',  # 06/2018-05/2022
            r'([A-Za-z]+\s+\d{4})\s*[-–—]\s*([A-Za-z]+\s+\d{4})',  # June 2018 - May 2022
            r'(\d{4})',  # Single year (graduation year)
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text)
            if match:
                if len(match.groups()) == 2:
                    entry['start_year'] = match.group(1)
                    entry['end_year'] = match.group(2)
                    entry['duration'] = f"{match.group(1)} - {match.group(2)}"
                elif len(match.groups()) == 1:
                    entry['end_year'] = match.group(1)
                    entry['duration'] = match.group(1)
                break
    
    def _extract_education_scores(self, text: str, entry: Dict[str, Any]):
        """Extract CGPA, percentage, or grades from education text."""
        text_lower = text.lower()
        
        # CGPA patterns
        cgpa_patterns = [
            r'cgpa[:\s]*(\d+\.?\d*)/(\d+\.?\d*)',  # CGPA: 8.5/10
            r'cgpa[:\s]*(\d+\.?\d*)',  # CGPA: 8.5
            r'gpa[:\s]*(\d+\.?\d*)/(\d+\.?\d*)',  # GPA: 3.8/4.0
            r'gpa[:\s]*(\d+\.?\d*)',  # GPA: 3.8
        ]
        
        for pattern in cgpa_patterns:
            match = re.search(pattern, text_lower)
            if match:
                if len(match.groups()) == 2:
                    entry['cgpa'] = f"{match.group(1)}/{match.group(2)}"
                else:
                    entry['cgpa'] = match.group(1)
                break
        
        # Percentage patterns
        percentage_patterns = [
            r'(\d+\.?\d*)%',  # 85.5%
            r'percentage[:\s]*(\d+\.?\d*)',  # Percentage: 85
        ]
        
        for pattern in percentage_patterns:
            match = re.search(pattern, text_lower)
            if match:
                entry['percentage'] = f"{match.group(1)}%"
                break
        
        # Grade patterns
        grade_patterns = [
            r'grade[:\s]*([A-F][+-]?)',  # Grade: A+
            r'\b([A-F][+-]?)\s+grade',  # A+ Grade
            r'first\s+class',  # First Class
            r'distinction',  # Distinction
        ]
        
        for pattern in grade_patterns:
            match = re.search(pattern, text_lower)
            if match:
                if 'first class' in text_lower:
                    entry['grade'] = 'First Class'
                elif 'distinction' in text_lower:
                    entry['grade'] = 'Distinction'
                else:
                    entry['grade'] = match.group(1).upper()
                break
    
    def extract_work_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract comprehensive work experience with detailed parsing."""
        work_entries = []
        
        lines = text.split('\n')
        in_experience_section = False
        current_entry = {}
        
        # Job title keywords for better detection
        job_title_indicators = [
            'engineer', 'developer', 'manager', 'analyst', 'specialist', 'coordinator',
            'director', 'senior', 'junior', 'lead', 'architect', 'consultant',
            'supervisor', 'administrator', 'technician', 'associate', 'assistant',
            'executive', 'officer', 'representative', 'advisor', 'researcher'
        ]
        
        # Company indicators
        company_indicators = [
            'inc', 'corp', 'ltd', 'llc', 'company', 'corporation', 'limited',
            'technologies', 'solutions', 'services', 'consulting', 'systems'
        ]
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check if entering experience section
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns['experience']):
                in_experience_section = True
                continue
            
            # Check if leaving experience section
            if in_experience_section:
                other_sections = ['education', 'skills', 'projects', 'certifications', 'achievements']
                if any(any(re.search(pattern, line_lower) for pattern in self.section_patterns[section]) 
                       for section in other_sections):
                    # Save current entry if exists
                    if current_entry:
                        work_entries.append(current_entry)
                        current_entry = {}
                    break
            
            if in_experience_section:
                # Check for job title patterns (usually first in experience entry)
                title_found = False
                for indicator in job_title_indicators:
                    if indicator in line_lower:
                        # Save previous entry if exists
                        if current_entry:
                            work_entries.append(current_entry)
                        
                        # Start new entry
                        current_entry = {
                            'job_title': line_clean,
                            'company': '',
                            'location': '',
                            'start_date': '',
                            'end_date': '',
                            'duration': '',
                            'employment_type': '',
                            'responsibilities': [],
                            'achievements': [],
                            'technologies': [],
                            'salary': '',
                            'team_size': ''
                        }
                        
                        # Extract dates from current line
                        self._extract_work_dates(line_clean, current_entry)
                        title_found = True
                        break
                
                # Look for company information
                if (not title_found and current_entry and 
                    not current_entry['company'] and
                    (any(indicator in line_lower for indicator in company_indicators) or
                     len(line_clean.split()) <= 4)):  # Short lines often company names
                    current_entry['company'] = line_clean
                    
                    # Extract location from company line
                    self._extract_location_from_line(line_clean, current_entry)
                
                # Look for bullet points (responsibilities/achievements)
                elif line_clean.startswith(('•', '-', '*', '▪', '◦')) or re.match(r'^\d+\.', line_clean):
                    if current_entry:
                        bullet_text = re.sub(r'^[•\-*▪◦\d\.]\s*', '', line_clean)
                        
                        # Categorize as achievement or responsibility
                        if self._is_achievement(bullet_text):
                            current_entry['achievements'].append(bullet_text)
                        else:
                            current_entry['responsibilities'].append(bullet_text)
                        
                        # Extract technologies from bullet points
                        tech_found = self._extract_technologies_from_text(bullet_text)
                        current_entry['technologies'].extend(tech_found)
                
                # Look for additional details
                elif current_entry:
                    # Check for employment type
                    employment_types = ['full-time', 'part-time', 'contract', 'freelance', 'internship', 'temporary']
                    for emp_type in employment_types:
                        if emp_type in line_lower and not current_entry['employment_type']:
                            current_entry['employment_type'] = emp_type.title()
                    
                    # Check for team size
                    team_match = re.search(r'team\s+of\s+(\d+)', line_lower)
                    if team_match and not current_entry['team_size']:
                        current_entry['team_size'] = f"{team_match.group(1)} members"
                    
                    # Check for salary information
                    salary_patterns = [
                        r'\$[\d,]+(?:k|,000)?',
                        r'salary[:\s]*\$?[\d,]+',
                        r'compensation[:\s]*\$?[\d,]+'
                    ]
                    for pattern in salary_patterns:
                        salary_match = re.search(pattern, line_lower)
                        if salary_match and not current_entry['salary']:
                            current_entry['salary'] = salary_match.group(0)
        
        # Save final entry
        if current_entry:
            work_entries.append(current_entry)
        
        # Clean up entries
        cleaned_entries = []
        for entry in work_entries:
            if entry.get('job_title') or entry.get('company'):
                # Remove duplicates from technologies
                if entry['technologies']:
                    entry['technologies'] = list(set(entry['technologies']))
                
                # Calculate duration if not present
                if not entry['duration'] and entry['start_date'] and entry['end_date']:
                    entry['duration'] = self._calculate_duration(entry['start_date'], entry['end_date'])
                
                cleaned_entries.append(entry)
        
        return cleaned_entries
    
    def extract_skills(self, text: str) -> Dict[str, Any]:
        """Extract comprehensive skills with categorization and proficiency detection."""
        skills_data = {
            'technical_skills': {
                'programming_languages': [],
                'web_development': [],
                'frameworks_libraries': [],
                'databases': [],
                'cloud_platforms': [],
                'devops_cicd': [],
                'version_control': [],
                'testing_qa': [],
                'mobile_development': [],
                'data_science_ai': [],
                'cybersecurity': [],
                'business_tools': []
            },
            'soft_skills': [],
            'languages': [],
            'proficiency_levels': {},
            'certifications_related': []
        }
        
        # First, extract from dedicated skills section
        skills_section_text = self._extract_skills_section(text)
        
        # Then scan entire document for skills
        all_text = text.lower()
        
        # Extract technical skills by category
        for category, category_data in self.tech_skills_database.items():
            found_skills = []
            
            # Get keywords and synonyms from the category data
            keywords = category_data.get('keywords', [])
            synonym_map = category_data.get('synonyms', {})
            
            for skill_name in keywords:
                # Check for skill and its synonyms
                all_variations = [skill_name]
                # Add synonyms that point to this skill
                for synonym, canonical in synonym_map.items():
                    if canonical == skill_name:
                        all_variations.append(synonym)
                
                for variation in all_variations:
                    # Pattern to match skill (whole word, case insensitive)
                    pattern = r'\b' + re.escape(variation.lower()) + r'\b'
                    if re.search(pattern, all_text):
                        # Find proficiency level if mentioned nearby
                        proficiency = self._extract_skill_proficiency(text, variation)
                        
                        skill_entry = {
                            'name': skill_name,
                            'matched_as': variation,
                            'proficiency': proficiency,
                            'years_experience': self._extract_skill_experience(text, variation)
                        }
                        
                        found_skills.append(skill_entry)
                        
                        # Store proficiency mapping
                        if proficiency:
                            skills_data['proficiency_levels'][skill_name] = proficiency
                        
                        break  # Don't match multiple synonyms for same skill
            
            skills_data['technical_skills'][category] = found_skills
        
        # Extract soft skills
        skills_data['soft_skills'] = self._extract_soft_skills(text)
        
        # Extract languages
        skills_data['languages'] = self._extract_languages(text)
        
        # Extract certifications mentioned in skills context
        skills_data['certifications_related'] = self._extract_skills_certifications(text)
        
        return skills_data
    
    def _extract_skills_section(self, text: str) -> str:
        """Extract text from dedicated skills section."""
        lines = text.split('\n')
        in_skills_section = False
        skills_lines = []
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check if entering skills section
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns['skills']):
                in_skills_section = True
                continue
            
            # Check if leaving skills section
            if in_skills_section:
                other_sections = ['education', 'experience', 'projects', 'certifications', 'achievements']
                if any(any(re.search(pattern, line_lower) for pattern in self.section_patterns[section]) 
                       for section in other_sections):
                    break
                
                skills_lines.append(line_clean)
        
        return ' '.join(skills_lines)

    def _extract_skill_proficiency(self, text: str, skill: str) -> str:
        """Extract proficiency level for a specific skill."""
        # Proficiency keywords
        proficiency_patterns = {
            'expert': ['expert', 'advanced', 'senior', 'lead', 'architect'],
            'proficient': ['proficient', 'intermediate', 'experienced', 'skilled'],
            'beginner': ['beginner', 'basic', 'novice', 'learning', 'familiar'],
            'years': [r'(\d+)\+?\s*years?', r'(\d+)\+?\s*yrs?']
        }
        
        # Look for proficiency mentions near the skill
        skill_pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        text_lower = text.lower()
        
        for match in re.finditer(skill_pattern, text_lower):
            start = max(0, match.start() - 100)
            end = min(len(text_lower), match.end() + 100)
            context = text_lower[start:end]
            
            # Check for proficiency levels
            for level, patterns in proficiency_patterns.items():
                if level == 'years':
                    for pattern in patterns:
                        years_match = re.search(pattern, context)
                        if years_match:
                            return f"{years_match.group(1)} years"
                else:
                    for pattern in patterns:
                        if pattern in context:
                            return level.title()
        
        return ""
    
    def _extract_skill_experience(self, text: str, skill: str) -> str:
        """Extract years of experience for a specific skill."""
        skill_pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        text_lower = text.lower()
        
        for match in re.finditer(skill_pattern, text_lower):
            start = max(0, match.start() - 150)
            end = min(len(text_lower), match.end() + 150)
            context = text_lower[start:end]
            
            # Look for years of experience
            experience_patterns = [
                r'(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)',
                r'(\d+)\+?\s*yrs?\s*(?:experience|exp)',
                r'experience[:\s]*(\d+)\+?\s*years?',
                r'(\d+)\+?\s*years?\s*in'
            ]
            
            for pattern in experience_patterns:
                exp_match = re.search(pattern, context)
                if exp_match:
                    return f"{exp_match.group(1)} years"
        
        return ""
    
    def _extract_soft_skills(self, text: str) -> List[Dict[str, str]]:
        """Extract soft skills with context."""
        soft_skills = []
        text_lower = text.lower()
        
        for skill in self.soft_skills:
            pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(pattern, text_lower):
                soft_skills.append({
                    'name': skill,
                    'context': self._get_skill_context(text, skill)
                })
        
        return soft_skills
    
    def _extract_languages(self, text: str) -> List[Dict[str, str]]:
        """Extract spoken languages with proficiency."""
        languages = []
        text_lower = text.lower()
        
        # Common languages
        language_list = [
            'English', 'Spanish', 'French', 'German', 'Italian', 'Portuguese',
            'Mandarin', 'Chinese', 'Japanese', 'Korean', 'Arabic', 'Russian',
            'Hindi', 'Bengali', 'Tamil', 'Telugu', 'Marathi', 'Gujarati',
            'Punjabi', 'Urdu', 'Malayalam', 'Kannada', 'Odia', 'Assamese'
        ]
        
        # Proficiency levels
        proficiency_keywords = {
            'native': ['native', 'mother tongue', 'first language'],
            'fluent': ['fluent', 'advanced', 'professional'],
            'conversational': ['conversational', 'intermediate', 'working'],
            'basic': ['basic', 'beginner', 'elementary']
        }
        
        for language in language_list:
            pattern = r'\b' + re.escape(language.lower()) + r'\b'
            if re.search(pattern, text_lower):
                # Find proficiency level
                proficiency = 'Not specified'
                
                # Look for proficiency near language mention
                lang_matches = list(re.finditer(pattern, text_lower))
                for match in lang_matches:
                    start = max(0, match.start() - 50)
                    end = min(len(text_lower), match.end() + 50)
                    context = text_lower[start:end]
                    
                    for level, keywords in proficiency_keywords.items():
                        if any(keyword in context for keyword in keywords):
                            proficiency = level.title()
                            break
                    
                    if proficiency != 'Not specified':
                        break
                
                languages.append({
                    'name': language,
                    'proficiency': proficiency
                })
        
        return languages
    
    def _extract_skills_certifications(self, text: str) -> List[str]:
        """Extract certification-related skills."""
        certifications = []
        text_lower = text.lower()
        
        # Common certification keywords
        cert_patterns = [
            r'certified\s+(\w+(?:\s+\w+){0,3})',
            r'(\w+(?:\s+\w+){0,2})\s+certified',
            r'(\w+(?:\s+\w+){0,2})\s+certification',
            r'aws\s+(\w+(?:\s+\w+){0,2})',
            r'microsoft\s+(\w+(?:\s+\w+){0,2})',
            r'google\s+(\w+(?:\s+\w+){0,2})',
            r'oracle\s+(\w+(?:\s+\w+){0,2})',
            r'cisco\s+(\w+(?:\s+\w+){0,2})',
            r'comptia\s+(\w+(?:\s+\w+){0,2})'
        ]
        
        for pattern in cert_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                if isinstance(match, tuple):
                    cert_name = ' '.join(match).strip()
                else:
                    cert_name = match.strip()
                
                if cert_name and len(cert_name) > 2:
                    certifications.append(cert_name.title())
        
        return list(set(certifications))  # Remove duplicates
    
    def _get_skill_context(self, text: str, skill: str) -> str:
        """Get context where skill is mentioned."""
        skill_pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        text_lower = text.lower()
        
        match = re.search(skill_pattern, text_lower)
        if match:
            start = max(0, match.start() - 50)
            end = min(len(text_lower), match.end() + 50)
            context = text[start:end].strip()
            return context
        
        return ""
    
    def _extract_from_skills_section(self, text: str) -> List[Dict]:
        """Extract skills from dedicated skills sections."""
        skills = []
        lines = text.split('\n')
        in_skills_section = False
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
            
            # Check if we're entering a skills section
            line_lower = line_clean.lower()
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns['skills']):
                in_skills_section = True
                continue
            
            # Check if we're leaving the skills section
            if in_skills_section:
                # Stop if we hit another major section
                if any(section in line_lower for section in ['experience', 'education', 'project', 'certification']):
                    in_skills_section = False
                    continue
                
                # Extract skills from the line
                if line_clean.startswith(('•', '-', '*', '◦')):
                    skill_text = line_clean[1:].strip()
                else:
                    skill_text = line_clean
                
                # Split by common separators
                skill_items = re.split(r'[,;|•·]', skill_text)
                for item in skill_items:
                    item_clean = item.strip()
                    if item_clean and len(item_clean.split()) <= 3:  # Reasonable skill name length
                        skills.append({
                            'name': item_clean,
                            'proficiency': 'Not specified'
                        })
        
        return skills
    
    def extract_projects(self, text: str) -> List[Dict[str, Any]]:
        """Extract comprehensive project information with detailed parsing."""
        projects = []
        
        lines = text.split('\n')
        in_project_section = False
        current_project = {}
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check if entering projects section
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns['projects']):
                in_project_section = True
                continue
            
            # Check if leaving projects section
            if in_project_section:
                other_sections = ['education', 'experience', 'skills', 'certifications', 'achievements']
                if any(any(re.search(pattern, line_lower) for pattern in self.section_patterns[section]) 
                       for section in other_sections):
                    # Save current project if exists
                    if current_project:
                        projects.append(current_project)
                        current_project = {}
                    break
            
            if in_project_section:
                # Check for project titles (usually standalone lines, not bullet points)
                if (not line_clean.startswith(('•', '-', '*', '▪', '◦')) and 
                    not re.match(r'^\d+\.', line_clean) and
                    len(line_clean.split()) <= 8 and
                    len(line_clean) > 5):
                    
                    # Save previous project if exists
                    if current_project:
                        projects.append(current_project)
                    
                    # Start new project
                    current_project = {
                        'name': line_clean,
                        'description': '',
                        'technologies': [],
                        'duration': '',
                        'start_date': '',
                        'end_date': '',
                        'status': '',
                        'role': '',
                        'team_size': '',
                        'links': {
                            'github': '',
                            'demo': '',
                            'website': ''
                        },
                        'highlights': [],
                        'challenges': '',
                        'outcomes': ''
                    }
                    
                    # Extract dates from project title line
                    self._extract_project_dates(line_clean, current_project)
                
                # Handle bullet points and descriptions
                elif line_clean.startswith(('•', '-', '*', '▪', '◦')) or re.match(r'^\d+\.', line_clean):
                    if current_project:
                        bullet_text = re.sub(r'^[•\-*▪◦\d\.]\s*', '', line_clean)
                        
                        # Categorize bullet point content
                        if self._is_project_outcome(bullet_text):
                            current_project['outcomes'] = bullet_text
                        elif self._is_project_challenge(bullet_text):
                            current_project['challenges'] = bullet_text
                        else:
                            current_project['highlights'].append(bullet_text)
                        
                        # Extract technologies from bullet points
                        tech_found = self._extract_technologies_from_text(bullet_text)
                        current_project['technologies'].extend(tech_found)
                
                # Handle description lines (not bullets, not titles)
                elif current_project and len(line_clean.split()) > 3:
                    # Extract links
                    links_found = self._extract_project_links(line_clean)
                    for link_type, url in links_found.items():
                        if url and not current_project['links'][link_type]:
                            current_project['links'][link_type] = url
                    
                    # If no links found, treat as description
                    if not any(links_found.values()):
                        if not current_project['description']:
                            current_project['description'] = line_clean
                        else:
                            current_project['description'] += ' ' + line_clean
                    
                    # Extract role information
                    role_keywords = ['role:', 'position:', 'as a', 'as an', 'worked as']
                    for keyword in role_keywords:
                        if keyword in line_lower and not current_project['role']:
                            role_text = line_clean.split(keyword, 1)[1].strip()
                            current_project['role'] = role_text.split('.')[0].strip()
                    
                    # Extract team size
                    team_match = re.search(r'team\s+of\s+(\d+)', line_lower)
                    if team_match and not current_project['team_size']:
                        current_project['team_size'] = f"{team_match.group(1)} members"
        
        # Save final project
        if current_project:
            projects.append(current_project)
        
        # Clean up projects
        cleaned_projects = []
        for project in projects:
            if project.get('name'):
                # Remove duplicate technologies
                if project['technologies']:
                    project['technologies'] = list(set(project['technologies']))
                
                # Set status based on dates
                if not project['status']:
                    if project['end_date'] and 'present' not in project['end_date'].lower():
                        project['status'] = 'Completed'
                    elif project['start_date']:
                        project['status'] = 'In Progress'
                    else:
                        project['status'] = 'Not Specified'
                
                cleaned_projects.append(project)
        
        return cleaned_projects
    
    def _extract_project_dates(self, text: str, project: Dict[str, Any]):
        """Extract project dates from text."""
        date_patterns = [
            r'(\w+\s+\d{4})\s*[-–—]\s*(\w+\s+\d{4}|\w+)',  # January 2020 - March 2022
            r'(\d{1,2}/\d{4})\s*[-–—]\s*(\d{1,2}/\d{4}|present|current)',  # 01/2020 - 03/2022
            r'(\d{4})\s*[-–—]\s*(\d{4}|present|current)',  # 2020 - 2022
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                project['start_date'] = match.group(1)
                end_date = match.group(2)
                if end_date.lower() in ['present', 'current']:
                    project['end_date'] = 'Present'
                else:
                    project['end_date'] = end_date
                project['duration'] = f"{match.group(1)} - {project['end_date']}"
                break
    
    def _extract_project_links(self, text: str) -> Dict[str, str]:
        """Extract project-related links."""
        links = {'github': '', 'demo': '', 'website': ''}
        
        # GitHub patterns
        github_patterns = [
            r'github\.com/[\w\-\.]+/[\w\-\.]+',
            r'git@github\.com:[\w\-\.]+/[\w\-\.]+\.git'
        ]
        
        # Demo/Live site patterns
        demo_patterns = [
            r'demo[:\s]*(?:https?://)?([^\s]+)',
            r'live[:\s]*(?:https?://)?([^\s]+)',
            r'preview[:\s]*(?:https?://)?([^\s]+)'
        ]
        
        # General URL patterns
        url_pattern = r'https?://[^\s]+'
        
        text_lower = text.lower()
        
        # Extract GitHub links
        for pattern in github_patterns:
            match = re.search(pattern, text)
            if match:
                links['github'] = match.group(0)
                break
        
        # Extract demo links
        for pattern in demo_patterns:
            match = re.search(pattern, text_lower)
            if match:
                links['demo'] = match.group(1)
                break
        
        # Extract general website links (if not GitHub or demo)
        if not links['github'] and not links['demo']:
            url_match = re.search(url_pattern, text)
            if url_match:
                links['website'] = url_match.group(0)
        
        return links
    
    def _is_project_outcome(self, text: str) -> bool:
        """Determine if text describes project outcomes."""
        outcome_indicators = [
            'result', 'outcome', 'achieved', 'delivered', 'success', 'impact',
            'increased', 'improved', 'reduced', 'saved', 'generated'
        ]
        
        text_lower = text.lower()
        return any(indicator in text_lower for indicator in outcome_indicators)
    
    def _is_project_challenge(self, text: str) -> bool:
        """Determine if text describes project challenges."""
        challenge_indicators = [
            'challenge', 'problem', 'issue', 'difficulty', 'obstacle',
            'solved', 'overcame', 'addressed', 'tackled'
        ]
        
        text_lower = text.lower()
        return any(indicator in text_lower for indicator in challenge_indicators)
    
    def calculate_parsing_confidence(self, extracted_data: Dict) -> float:
        """Calculate confidence score based on successfully extracted fields."""
        total_fields = 8  # contact, education, experience, skills, etc.
        filled_fields = 0
        
        # Check contact_info with new structure
        contact_info = extracted_data.get('contact_info', {})
        if (contact_info and 
            (contact_info.get('name', {}).get('full') or 
             contact_info.get('emails') or 
             contact_info.get('phones') or 
             any(contact_info.get('location', {}).values()) or
             any(contact_info.get('links', {}).values()))):
            filled_fields += 1
            
        if extracted_data.get('education'):
            filled_fields += 1
        if extracted_data.get('work_experience'):
            filled_fields += 1
        if extracted_data.get('skills'):
            filled_fields += 1
        if extracted_data.get('projects'):
            filled_fields += 1
        if extracted_data.get('certifications'):
            filled_fields += 1
        if extracted_data.get('achievements'):
            filled_fields += 1
        if extracted_data.get('career_summary'):
            filled_fields += 1
        
        return (filled_fields / total_fields) * 100
    
    def extract_certifications(self, text: str) -> List[Dict[str, Any]]:
        """Extract comprehensive certifications and professional credentials."""
        certifications = []
        
        lines = text.split('\n')
        in_cert_section = False
        current_cert = {}
        
        # Enhanced certification keywords
        cert_keywords = {
            'aws': ['aws certified', 'amazon web services', 'cloud practitioner', 'solutions architect', 'developer associate'],
            'microsoft': ['microsoft certified', 'azure', 'mcse', 'mcsa', 'office specialist'],
            'google': ['google cloud', 'gcp certified', 'google certified', 'associate cloud engineer'],
            'cisco': ['cisco certified', 'ccna', 'ccnp', 'ccie', 'network associate'],
            'oracle': ['oracle certified', 'ocp', 'oca', 'database administrator'],
            'comptia': ['comptia', 'security+', 'network+', 'a+', 'linux+'],
            'project_management': ['pmp', 'capm', 'prince2', 'agile certified', 'scrum master', 'csm'],
            'security': ['cissp', 'ceh', 'cism', 'cisa', 'certified ethical hacker'],
            'data': ['tableau certified', 'power bi', 'certified analytics professional'],
            'general': ['certified', 'certification', 'certificate', 'credential', 'license']
        }
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check if entering certifications section
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns['certifications']):
                in_cert_section = True
                continue
            
            # Check if leaving certifications section
            if in_cert_section:
                other_sections = ['education', 'experience', 'skills', 'projects', 'achievements']
                if any(any(re.search(pattern, line_lower) for pattern in self.section_patterns[section]) 
                       for section in other_sections):
                    # Save current certification if exists
                    if current_cert:
                        certifications.append(current_cert)
                        current_cert = {}
                    break
            
            # Look for certifications throughout the document
            cert_found = False
            for category, keywords in cert_keywords.items():
                for keyword in keywords:
                    if keyword in line_lower:
                        # Save previous certification if exists
                        if current_cert:
                            certifications.append(current_cert)
                        
                        # Start new certification
                        current_cert = {
                            'name': line_clean,
                            'category': category.replace('_', ' ').title(),
                            'issuer': '',
                            'date_obtained': '',
                            'expiry_date': '',
                            'credential_id': '',
                            'verification_url': '',
                            'status': 'Active',
                            'score': '',
                            'level': ''
                        }
                        
                        # Extract details from current line
                        self._extract_cert_details(line_clean, current_cert)
                        
                        # Extract issuer based on category
                        current_cert['issuer'] = self._determine_cert_issuer(category, line_clean)
                        
                        cert_found = True
                        break
                
                if cert_found:
                    break
            
            # If in certification section or found cert, look for additional details
            if (in_cert_section or current_cert) and not cert_found:
                if current_cert:
                    # Extract additional details from nearby lines
                    self._extract_cert_additional_details(line_clean, current_cert)
        
        # Save final certification
        if current_cert:
            certifications.append(current_cert)
        
        # Clean up and deduplicate
        cleaned_certs = []
        seen_names = set()
        
        for cert in certifications:
            if cert.get('name') and cert['name'] not in seen_names:
                seen_names.add(cert['name'])
                
                # Determine level if not set
                if not cert['level']:
                    cert['level'] = self._determine_cert_level(cert['name'])
                
                # Set status based on expiry
                if cert['expiry_date'] and self._is_cert_expired(cert['expiry_date']):
                    cert['status'] = 'Expired'
                
                cleaned_certs.append(cert)
        
        return cleaned_certs
    
    def _extract_cert_details(self, text: str, cert: Dict[str, Any]):
        """Extract certification details from text."""
        text_lower = text.lower()
        
        # Extract dates
        date_patterns = [
            r'(\d{1,2}/\d{4})',  # MM/YYYY
            r'(\w+\s+\d{4})',  # Month YYYY
            r'(\d{4})',  # YYYY
            r'valid until[:\s]*(\d{1,2}/\d{4}|\w+\s+\d{4})',  # Expiry
            r'expires[:\s]*(\d{1,2}/\d{4}|\w+\s+\d{4})',  # Expiry
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if not cert['date_obtained']:
                    cert['date_obtained'] = match
                elif not cert['expiry_date'] and ('valid' in text_lower or 'expires' in text_lower):
                    cert['expiry_date'] = match
        
        # Extract credential ID
        id_patterns = [
            r'id[:\s]*([A-Za-z0-9\-]+)',
            r'credential[:\s]*([A-Za-z0-9\-]+)',
            r'license[:\s]*([A-Za-z0-9\-]+)',
            r'cert[:\s]*([A-Za-z0-9\-]+)'
        ]
        
        for pattern in id_patterns:
            id_match = re.search(pattern, text_lower)
            if id_match:
                cert['credential_id'] = id_match.group(1)
                break
        
        # Extract score
        score_patterns = [
            r'score[:\s]*(\d+(?:\.\d+)?)',
            r'(\d+)%',
            r'(\d+)/(\d+)'
        ]
        
        for pattern in score_patterns:
            score_match = re.search(pattern, text_lower)
            if score_match:
                if len(score_match.groups()) == 2:  # Fraction format
                    cert['score'] = f"{score_match.group(1)}/{score_match.group(2)}"
                else:
                    cert['score'] = score_match.group(1)
                break
        
        # Extract verification URL
        url_pattern = r'https?://[^\s]+'
        url_match = re.search(url_pattern, text)
        if url_match:
            cert['verification_url'] = url_match.group(0)
    
    def _extract_cert_additional_details(self, text: str, cert: Dict[str, Any]):
        """Extract additional certification details from nearby text."""
        text_lower = text.lower()
        
        # Extract missing dates
        if not cert['date_obtained'] or not cert['expiry_date']:
            date_patterns = [
                r'(\d{1,2}/\d{4})',
                r'(\w+\s+\d{4})',
                r'(\d{4})'
            ]
            
            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    if not cert['date_obtained']:
                        cert['date_obtained'] = match
                    elif not cert['expiry_date']:
                        cert['expiry_date'] = match
        
        # Extract missing credential ID
        if not cert['credential_id']:
            id_patterns = [
                r'(?:id|credential|license|cert)[:\s]*([A-Za-z0-9\-]+)',
                r'([A-Za-z0-9\-]{8,})'  # Long alphanumeric strings
            ]
            
            for pattern in id_patterns:
                id_match = re.search(pattern, text_lower)
                if id_match:
                    cert['credential_id'] = id_match.group(1)
                    break
        
        # Extract verification URL if missing
        if not cert['verification_url']:
            url_pattern = r'https?://[^\s]+'
            url_match = re.search(url_pattern, text)
            if url_match:
                cert['verification_url'] = url_match.group(0)
    
    def _determine_cert_issuer(self, category: str, text: str) -> str:
        """Determine certification issuer based on category and text."""
        issuer_map = {
            'aws': 'Amazon Web Services',
            'microsoft': 'Microsoft',
            'google': 'Google Cloud',
            'cisco': 'Cisco',
            'oracle': 'Oracle',
            'comptia': 'CompTIA',
            'project_management': self._extract_pm_issuer(text),
            'security': self._extract_security_issuer(text),
            'data': self._extract_data_issuer(text),
        }
        
        return issuer_map.get(category, 'Unknown')
    
    def _extract_pm_issuer(self, text: str) -> str:
        """Extract project management certification issuer."""
        text_lower = text.lower()
        if 'pmi' in text_lower or 'pmp' in text_lower:
            return 'Project Management Institute (PMI)'
        elif 'scrum' in text_lower:
            return 'Scrum Alliance'
        elif 'prince2' in text_lower:
            return 'AXELOS'
        return 'Unknown'
    
    def _extract_security_issuer(self, text: str) -> str:
        """Extract security certification issuer."""
        text_lower = text.lower()
        if 'isc2' in text_lower or 'cissp' in text_lower:
            return '(ISC)² - International Information System Security Certification Consortium'
        elif 'ec-council' in text_lower or 'ceh' in text_lower:
            return 'EC-Council'
        return 'Unknown'
    
    def _extract_data_issuer(self, text: str) -> str:
        """Extract data certification issuer."""
        text_lower = text.lower()
        if 'tableau' in text_lower:
            return 'Tableau'
        elif 'microsoft' in text_lower and 'power bi' in text_lower:
            return 'Microsoft'
        return 'Unknown'
    
    def _determine_cert_level(self, cert_name: str) -> str:
        """Determine certification level."""
        cert_name_lower = cert_name.lower()
        
        if any(keyword in cert_name_lower for keyword in ['associate', 'foundation', 'fundamentals']):
            return 'Associate'
        elif any(keyword in cert_name_lower for keyword in ['professional', 'specialist']):
            return 'Professional'
        elif any(keyword in cert_name_lower for keyword in ['expert', 'architect', 'master']):
            return 'Expert'
        elif any(keyword in cert_name_lower for keyword in ['advanced', 'senior']):
            return 'Advanced'
        
        return 'Standard'
    
    def _is_cert_expired(self, expiry_date: str) -> bool:
        """Check if certification is expired."""
        from datetime import datetime
        try:
            # Simple check - could be enhanced with proper date parsing
            current_year = datetime.now().year
            if expiry_date.isdigit() and int(expiry_date) < current_year:
                return True
            return False
        except:
            return False

    def extract_achievements(self, text: str) -> List[Dict[str, Any]]:
        """Extract comprehensive achievements, awards, and accomplishments."""
        achievements = []
        
        lines = text.split('\n')
        in_achievement_section = False
        current_achievement = {}
        
        # Achievement categories and keywords
        achievement_categories = {
            'awards': ['award', 'awarded', 'winner', 'champion', 'first place', 'second place', 'third place'],
            'recognition': ['recognition', 'recognized', 'honored', 'distinguished', 'outstanding'],
            'publications': ['published', 'publication', 'paper', 'article', 'journal', 'conference'],
            'patents': ['patent', 'invention', 'intellectual property'],
            'certifications': ['certified', 'certification', 'accredited'],
            'academic': ['dean\'s list', 'honor roll', 'magna cum laude', 'summa cum laude', 'valedictorian', 'scholarship'],
            'professional': ['employee of the month', 'top performer', 'exceeded targets', 'promotion'],
            'leadership': ['elected', 'appointed', 'led', 'founded', 'established', 'organized'],
            'media': ['featured', 'interviewed', 'quoted', 'mentioned', 'appeared'],
            'competitions': ['competition', 'contest', 'hackathon', 'tournament', 'championship']
        }
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            if not line_clean:
                continue
            
            line_lower = line_clean.lower()
            
            # Check if entering achievements section
            if any(re.search(pattern, line_lower) for pattern in self.section_patterns.get('achievements', [])):
                in_achievement_section = True
                continue
            
            # Check if leaving achievements section
            if in_achievement_section:
                other_sections = ['education', 'experience', 'skills', 'projects', 'certifications']
                if any(any(re.search(pattern, line_lower) for pattern in self.section_patterns.get(section, [])) 
                       for section in other_sections):
                    # Save current achievement if exists
                    if current_achievement:
                        achievements.append(current_achievement)
                        current_achievement = {}
                    break
            
            # Look for achievement indicators throughout the document
            achievement_found = False
            for category, keywords in achievement_categories.items():
                for keyword in keywords:
                    if keyword in line_lower:
                        # Save previous achievement if exists
                        if current_achievement:
                            achievements.append(current_achievement)
                        
                        # Start new achievement
                        current_achievement = {
                            'title': line_clean,
                            'category': category.title(),
                            'description': '',
                            'date': '',
                            'issuer': '',
                            'location': '',
                            'significance': '',
                            'impact': '',
                            'verification_url': '',
                            'associated_with': ''  # Company, university, organization
                        }
                        
                        # Extract details from current line
                        self._extract_achievement_details(line_clean, current_achievement)
                        achievement_found = True
                        break
                
                if achievement_found:
                    break
            
            # If in achievement section or found achievement, look for additional details
            if (in_achievement_section or current_achievement) and not achievement_found:
                if current_achievement:
                    # Extract additional details from nearby lines
                    self._extract_achievement_additional_details(line_clean, current_achievement)
        
        # Save final achievement
        if current_achievement:
            achievements.append(current_achievement)
        
        # Clean up and enhance achievements
        cleaned_achievements = []
        for achievement in achievements:
            if achievement.get('title'):
                # Determine significance level
                if not achievement['significance']:
                    achievement['significance'] = self._determine_achievement_significance(achievement['title'])
                
                # Extract impact if not set
                if not achievement['impact']:
                    achievement['impact'] = self._extract_achievement_impact(achievement['title'])
                
                cleaned_achievements.append(achievement)
        
        return cleaned_achievements
    
    def _extract_achievement_details(self, text: str, achievement: Dict[str, Any]):
        """Extract achievement details from text."""
        text_lower = text.lower()
        
        # Extract dates
        date_patterns = [
            r'(\d{1,2}/\d{4})',  # MM/YYYY
            r'(\w+\s+\d{4})',  # Month YYYY
            r'(\d{4})',  # YYYY
        ]
        
        for pattern in date_patterns:
            date_match = re.search(pattern, text)
            if date_match:
                achievement['date'] = date_match.group(0)
                break
        
        # Extract issuer/organization
        org_patterns = [
            r'by\s+([A-Z][a-zA-Z\s&.]+)',  # "by Organization"
            r'from\s+([A-Z][a-zA-Z\s&.]+)',  # "from Organization"
            r'at\s+([A-Z][a-zA-Z\s&.]+)',  # "at Organization"
        ]
        
        for pattern in org_patterns:
            org_match = re.search(pattern, text)
            if org_match:
                potential_org = org_match.group(1).strip()
                # Filter out common words that aren't organizations
                if len(potential_org.split()) <= 5 and not any(word in potential_org.lower() 
                    for word in ['the', 'and', 'of', 'for', 'in', 'with']):
                    achievement['issuer'] = potential_org
                break
        
        # Extract location
        location_patterns = [
            r'([A-Z][a-z]+,\s*[A-Z]{2})',  # City, ST
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+)',  # City, State
        ]
        
        for pattern in location_patterns:
            location_match = re.search(pattern, text)
            if location_match:
                achievement['location'] = location_match.group(1)
                break
        
        # Extract verification URL
        url_pattern = r'https?://[^\s]+'
        url_match = re.search(url_pattern, text)
        if url_match:
            achievement['verification_url'] = url_match.group(0)
    
    def _extract_achievement_additional_details(self, text: str, achievement: Dict[str, Any]):
        """Extract additional achievement details from nearby text."""
        text_lower = text.lower()
        
        # If it's a description line (longer text, not bullet points)
        if (len(text.split()) > 5 and 
            not text.startswith(('•', '-', '*', '▪', '◦')) and
            not achievement['description']):
            achievement['description'] = text
        
        # Extract missing details
        if not achievement['date']:
            date_patterns = [r'(\d{1,2}/\d{4})', r'(\w+\s+\d{4})', r'(\d{4})']
            for pattern in date_patterns:
                date_match = re.search(pattern, text)
                if date_match:
                    achievement['date'] = date_match.group(0)
                    break
        
        if not achievement['issuer']:
            org_patterns = [
                r'by\s+([A-Z][a-zA-Z\s&.]+)',
                r'from\s+([A-Z][a-zA-Z\s&.]+)',
                r'at\s+([A-Z][a-zA-Z\s&.]+)'
            ]
            for pattern in org_patterns:
                org_match = re.search(pattern, text)
                if org_match:
                    achievement['issuer'] = org_match.group(1).strip()
                    break
        
        if not achievement['verification_url']:
            url_pattern = r'https?://[^\s]+'
            url_match = re.search(url_pattern, text)
            if url_match:
                achievement['verification_url'] = url_match.group(0)
    
    def _determine_achievement_significance(self, title: str) -> str:
        """Determine the significance level of an achievement."""
        title_lower = title.lower()
        
        high_significance = ['first place', 'winner', 'champion', 'top', 'best', 'outstanding', 'distinguished']
        medium_significance = ['second place', 'runner-up', 'finalist', 'recognition', 'honored']
        
        if any(keyword in title_lower for keyword in high_significance):
            return 'High'
        elif any(keyword in title_lower for keyword in medium_significance):
            return 'Medium'
        else:
            return 'Standard'
    
    def _extract_achievement_impact(self, title: str) -> str:
        """Extract impact description from achievement title."""
        title_lower = title.lower()
        
        # Look for quantifiable impact
        impact_patterns = [
            r'(\d+)%',  # Percentage
            r'(\d+)\s*(?:million|thousand|k)',  # Numbers
            r'saved\s*\$?(\d+(?:,\d+)*)',  # Money saved
            r'increased\s*(\d+)%',  # Percentage increase
            r'improved\s*(\d+)%',  # Percentage improvement
        ]
        
        for pattern in impact_patterns:
            impact_match = re.search(pattern, title_lower)
            if impact_match:
                return f"Quantified impact: {impact_match.group(0)}"
        
        # Look for qualitative impact
        if any(keyword in title_lower for keyword in ['increased', 'improved', 'enhanced', 'optimized']):
            return 'Performance improvement'
        elif any(keyword in title_lower for keyword in ['led', 'managed', 'directed']):
            return 'Leadership achievement'
        elif any(keyword in title_lower for keyword in ['published', 'presented', 'spoke']):
            return 'Knowledge sharing'
        
        return 'Professional recognition'

    def extract_courses(self, text: str) -> List[Dict]:
        """Extract relevant courses and training."""
        courses = []
        
        course_keywords = [
            'course', 'training', 'workshop', 'seminar', 'bootcamp', 'program',
            'coursera', 'udemy', 'edx', 'pluralsight', 'linkedin learning'
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower()
            
            if any(keyword in line_lower for keyword in course_keywords):
                course_entry = {
                    'name': line.strip(),
                    'provider': '',
                    'completion_date': ''
                }
                
                # Look for provider and date in nearby lines
                search_range = range(max(0, i-1), min(len(lines), i+2))
                for j in search_range:
                    if j != i:
                        search_line = lines[j]
                        
                        # Extract completion date
                        date_pattern = r'(\d{1,2}/\d{4}|\d{4}|[A-Za-z]+\s+\d{4})'
                        date_match = re.search(date_pattern, search_line)
                        if date_match:
                            course_entry['completion_date'] = date_match.group(0)
                
                courses.append(course_entry)
        
        return courses

    def extract_languages(self, text: str) -> List[Dict]:
        """Extract language skills and proficiency levels."""
        languages = []
        
        # Common languages
        language_names = [
            'english', 'spanish', 'french', 'german', 'italian', 'portuguese',
            'chinese', 'mandarin', 'japanese', 'korean', 'arabic', 'hindi',
            'russian', 'dutch', 'swedish', 'norwegian', 'polish'
        ]
        
        proficiency_levels = [
            'native', 'fluent', 'advanced', 'intermediate', 'basic', 'beginner',
            'conversational', 'professional', 'limited'
        ]
        
        lines = text.split('\n')
        for line in lines:
            line_lower = line.lower()
            
            # Check if line contains language keywords
            if 'language' in line_lower or any(lang in line_lower for lang in language_names):
                for lang in language_names:
                    if lang in line_lower:
                        lang_entry = {
                            'language': lang.title(),
                            'proficiency': 'Not specified'
                        }
                        
                        # Look for proficiency level
                        for level in proficiency_levels:
                            if level in line_lower:
                                lang_entry['proficiency'] = level.title()
                                break
                        
                        languages.append(lang_entry)
        
        return languages

    def extract_career_summary(self, text: str) -> str:
        """Extract career summary or objective."""
        summary_keywords = [
            'summary', 'objective', 'profile', 'overview', 'about', 'bio',
            'professional summary', 'career objective', 'personal statement'
        ]
        
        lines = text.split('\n')
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if line is a summary header
            if any(keyword in line_lower for keyword in summary_keywords) and len(line.split()) <= 5:
                # Extract the next few lines as summary
                summary_lines = []
                for j in range(i + 1, min(len(lines), i + 6)):
                    next_line = lines[j].strip()
                    if next_line and not any(keyword in next_line.lower() for keyword in ['education', 'experience', 'skills']):
                        summary_lines.append(next_line)
                    else:
                        break
                
                if summary_lines:
                    return ' '.join(summary_lines)
        
        # If no explicit summary found, take first few substantial lines
        substantial_lines = []
        for line in lines[:10]:
            line_clean = line.strip()
            if len(line_clean.split()) > 5 and not any(pattern in line_clean.lower() for pattern in ['phone', 'email', '@', 'linkedin']):
                substantial_lines.append(line_clean)
                if len(' '.join(substantial_lines).split()) > 50:
                    break
        
        return ' '.join(substantial_lines)

    def extract_custom_sections(self, text: str) -> Dict:
        """Extract any custom sections not covered by standard categories."""
        custom_sections = {}
        
        # Common custom section headers
        custom_headers = [
            'volunteer', 'volunteering', 'publications', 'patents', 'references',
            'interests', 'hobbies', 'activities', 'memberships', 'affiliations'
        ]
        
        lines = text.split('\n')
        current_section = None
        section_content = []
        
        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue
            
            # Check if line is a custom section header
            line_lower = line_clean.lower()
            for header in custom_headers:
                if header in line_lower and len(line_clean.split()) <= 3:
                    # Save previous section
                    if current_section and section_content:
                        custom_sections[current_section] = '\n'.join(section_content)
                    
                    # Start new section
                    current_section = header.title()
                    section_content = []
                    break
            else:
                # Add content to current section
                if current_section:
                    section_content.append(line_clean)
        
        # Save final section
        if current_section and section_content:
            custom_sections[current_section] = '\n'.join(section_content)
        
        return custom_sections

    def parse_resume(self, text: str) -> Dict:
        """Main parsing function that coordinates all comprehensive extraction methods."""
        try:
            # Log the extracted text for debugging
            self.logger.info(f"Parsing resume text (length: {len(text)} chars)")
            self.logger.debug(f"Text preview: {text[:500]}...")
            
            # Extract all resume components with comprehensive parsing
            extracted_data = {
                'contact_info': self.extract_personal_info(text),  # Updated to use comprehensive personal info
                'career_summary': self.extract_summary(text),  # Updated to use comprehensive summary
                'education': self.extract_education(text),  # Already comprehensive
                'work_experience': self.extract_work_experience(text),  # Already comprehensive
                'skills': self.extract_skills(text),  # Already comprehensive
                'projects': self.extract_projects(text),  # Already comprehensive
                'certifications': self.extract_certifications(text),  # Already comprehensive
                'achievements': self.extract_achievements(text),  # Already comprehensive
                'courses': self.extract_courses(text),
                'languages': self.extract_languages(text),
                'custom_sections': self.extract_custom_sections(text)
            }
            
            # Calculate parsing confidence based on comprehensive data
            confidence = self.calculate_parsing_confidence(extracted_data)
            extracted_data['parsing_confidence'] = confidence
            extracted_data['parsed_at'] = datetime.utcnow()
            extracted_data['raw_text'] = text[:1000] + "..." if len(text) > 1000 else text
            
            # Add metadata about extraction methods used
            extracted_data['extraction_metadata'] = {
                'total_sections_found': len([k for k, v in extracted_data.items() if v and k not in ['parsing_confidence', 'parsed_at', 'raw_text']]),
                'comprehensive_parsing': True,
                'extraction_methods': [
                    'personal_info_structured',
                    'education_detailed',
                    'work_experience_comprehensive',
                    'skills_categorized_with_proficiency',
                    'projects_detailed',
                    'certifications_comprehensive',
                    'achievements_categorized',
                    'summary_enhanced'
                ]
            }
            
            self.logger.info(f"Comprehensive parsing completed with confidence: {confidence}%")
            self.logger.info(f"Sections extracted: {extracted_data['extraction_metadata']['total_sections_found']}")
            
            return extracted_data
            
        except Exception as e:
            self.logger.error(f"Error during comprehensive resume parsing: {e}")
            # Return comprehensive structure on parsing failure
            return {
                'contact_info': {
                    'name': {'first': '', 'middle': '', 'last': '', 'full': ''},
                    'emails': [], 'phones': [], 'location': {}, 'links': {}
                },
                'career_summary': '',
                'education': [],
                'work_experience': [],
                'skills': {
                    'technical_skills': {}, 'soft_skills': [], 'languages': [],
                    'proficiency_levels': {}, 'certifications_related': []
                },
                'projects': [],
                'certifications': [],
                'achievements': [],
                'courses': [],
                'languages': [],
                'custom_sections': {},
                'parsing_confidence': 0.0,
                'parsed_at': datetime.utcnow(),
                'raw_text': text[:500] + "..." if len(text) > 500 else text,
                'error': str(e),
                'extraction_metadata': {
                    'total_sections_found': 0,
                    'comprehensive_parsing': False,
                    'error_occurred': True
                }
            }
    
    def _extract_work_dates(self, text: str, entry: Dict[str, Any]):
        """Extract work experience dates."""
        # Date patterns for work experience
        date_patterns = [
            r'(\w+\s+\d{4})\s*[-–—]\s*(\w+\s+\d{4}|\w+)',  # January 2020 - March 2022
            r'(\d{1,2}/\d{4})\s*[-–—]\s*(\d{1,2}/\d{4}|present|current)',  # 01/2020 - 03/2022
            r'(\d{4})\s*[-–—]\s*(\d{4}|present|current)',  # 2020 - 2022
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                entry['start_date'] = match.group(1)
                end_date = match.group(2)
                if end_date.lower() in ['present', 'current']:
                    entry['end_date'] = 'Present'
                else:
                    entry['end_date'] = end_date
                entry['duration'] = f"{match.group(1)} - {entry['end_date']}"
                break
    
    def _extract_location_from_line(self, text: str, entry: Dict[str, Any]):
        """Extract location information from text line."""
        # Common location patterns
        location_patterns = [
            r'([A-Z][a-z]+,\s*[A-Z]{2})',  # City, ST
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+)',  # City, State
            r'([A-Z][a-z]+\s+[A-Z][a-z]+,\s*[A-Z]{2})',  # New York, NY
            r'(Remote)',  # Remote work
            r'([A-Z][a-z]+,\s*[A-Z][a-z]+,\s*[A-Z][a-z]+)',  # City, State, Country
        ]
        
        for pattern in location_patterns:
            match = re.search(pattern, text)
            if match:
                entry['location'] = match.group(1)
                break
    
    def _is_achievement(self, text: str) -> bool:
        """Determine if text describes an achievement vs regular responsibility."""
        achievement_indicators = [
            'increased', 'improved', 'reduced', 'achieved', 'delivered', 'led',
            'managed', 'saved', 'generated', 'boosted', 'enhanced', 'optimized',
            'implemented', 'launched', 'created', 'developed', 'awarded',
            'recognized', 'promoted', 'exceeded', 'surpassed', '%', 'million',
            'thousand', 'decrease', 'increase', 'growth', 'revenue', 'profit'
        ]
        
        text_lower = text.lower()
        return any(indicator in text_lower for indicator in achievement_indicators)
    
    def _extract_technologies_from_text(self, text: str) -> List[str]:
        """Extract technology mentions from text."""
        technologies = []
        text_lower = text.lower()
        
        # Check against our tech skills database
        for category, category_data in self.tech_skills_database.items():
            keywords = category_data.get('keywords', [])
            synonym_map = category_data.get('synonyms', {})
            
            for skill_name in keywords:
                # Check for skill and its synonyms
                all_variations = [skill_name]
                # Add synonyms that point to this skill
                for synonym, canonical in synonym_map.items():
                    if canonical == skill_name:
                        all_variations.append(synonym)
                
                for variation in all_variations:
                    pattern = r'\b' + re.escape(variation.lower()) + r'\b'
                    if re.search(pattern, text_lower):
                        technologies.append(skill_name)
                        break
        
        return technologies
    
    def _calculate_duration(self, start_date: str, end_date: str) -> str:
        """Calculate duration between two dates."""
        try:
            # Simple duration calculation (could be enhanced with actual date parsing)
            if 'present' in end_date.lower():
                return f"{start_date} - Present"
            else:
                return f"{start_date} - {end_date}"
        except:
            return f"{start_date} - {end_date}"