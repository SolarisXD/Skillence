"""
Course PDF / DOCX Parser
=========================
Extracts structured syllabus content from individual course PDF/DOCX files.

VIT-style course documents follow a common structure:
  - Course Code, Name, Type, Credits
  - Prerequisites
  - Course Objectives
  - Course / Expected Outcomes
  - Unit Content table (Unit No | Content | Hours | SOs)
  - Mode of Teaching / Evaluation
  - Text Books / Reference Books
  - List of Experiments (optional)

This parser extracts the **syllabus-relevant sections** (unit content,
experiments, objectives) so they can be sent to Gemini for accurate
skill extraction.
"""

from __future__ import annotations

import re
import logging
from typing import Dict, Optional

logger = logging.getLogger(__name__)

# ──────────────────────────────────────────────────────────────────
# Section boundary patterns
# ──────────────────────────────────────────────────────────────────

# Markers that signal the START of unit content
_UNIT_START_RE = re.compile(
    r"^unit\s*(?:no|content|$)|^module\s*(?:no\.?|\d|description|$)",
    re.IGNORECASE,
)

# Header noise lines that appear between the table header row and actual content
# e.g. "No.", "Module Description", "No.of", "Hours", "SO"
_TABLE_HEADER_NOISE_RE = re.compile(
    r"^(?:"
    r"module\s*(?:no\.?|description|content)?|"
    r"unit\s*(?:no\.?|description|content)?|"
    r"no\.?\s*(?:of\s*(?:hours?)?)?|"
    r"hrs\.?|hours?|"
    r"s\.?o\.?|slo|sl\.?\s*no\.?|"
    r"description|content|"
    r"topics(?:\s+to\s+be\s+discussed)?|"
    r"syllabus\s*content"
    r")$",
    re.IGNORECASE,
)

# Some PDFs use "CO" (Course Outcome) sections instead of "Unit" sections
# e.g. "CO1", "CO2", etc. with topics listed under each
_CO_SECTION_RE = re.compile(r"^CO\d+$")

# Header row that precedes CO-based content  (e.g. "CO  Topics to be discussed  Hrs.")
_CO_TABLE_HEADER_RE = re.compile(r"^CO\s+(Topics|Syllabus|Content)", re.IGNORECASE)

# Markers that signal the END of unit content
_UNIT_END_KEYWORDS = [
    "mode of teaching",
    "mode of evaluation",
    "text book",
    "textbook",
    "reference book",
    "recommended book",
    "recommendation by the board",
    "approval by academic",
    "compiled by",
    "list of suggested experiment",
    "indicative list of experiment",
    "list of experiment",
    "laboratory experiments",
    "lab experiments",
]

# Markers for experiments section
_EXP_START_KEYWORDS = [
    "list of suggested experiment",
    "indicative list of experiment",
    "list of experiment",
    "laboratory experiments",
    "lab experiments",
]

_EXP_END_KEYWORDS = [
    "recommendation by the board",
    "approval by academic",
    "compiled by",
    "text book",
    "textbook",
    "reference book",
]

# Markers for objectives section
_OBJ_START_KEYWORDS = [
    "course objectives",
    "objectives:",
    "course  objectives",
]

_OBJ_END_KEYWORDS = [
    "course outcomes",
    "expected outcomes",
    "course  outcomes",
    "student outcomes",
]


def _line_starts_with_any(line: str, keywords: list[str]) -> bool:
    """Check if line starts with any keyword (case-insensitive)."""
    low = line.strip().lower()
    return any(low.startswith(kw) for kw in keywords)


def _clean_text(text: str) -> str:
    """Normalise whitespace and remove junk characters."""
    # Replace common PDF artefacts
    text = text.replace("\uf0b7", "-")  # bullet
    text = text.replace("\u30fb", "-")  # middle dot
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    # Collapse multiple spaces/newlines
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ──────────────────────────────────────────────────────────────────
# Extract sections from raw text
# ──────────────────────────────────────────────────────────────────

def _extract_unit_content(text: str) -> str:
    """Extract the unit/module content table section.

    Handles two common VIT formats:
      - Unit-based: "Unit No | Unit Content | Hours | SOs"
      - CO-based:   "CO | Topics to be discussed | Hrs."
    """
    lines = text.split("\n")

    # --- Try Unit-based / Module-based format ---
    result_lines: list[str] = []
    capturing = False

    for line in lines:
        stripped = line.strip()
        low = stripped.lower()

        if not capturing and _UNIT_START_RE.match(stripped):
            capturing = True
            continue

        if capturing:
            if _line_starts_with_any(stripped, _UNIT_END_KEYWORDS):
                break
            # Skip table header noise (No., Description, Hours, SO, etc.)
            if _TABLE_HEADER_NOISE_RE.match(stripped):
                continue
            if stripped.isdigit() or re.match(r"^\d+(\.\d+)?$", stripped):
                continue
            if re.match(r"^[a-l](?:\s*,\s*[a-l])*$", stripped):
                continue
            if not stripped:
                continue
            if low.startswith("total"):
                break
            if low.startswith("guest lecture"):
                break
            result_lines.append(stripped)

    if result_lines:
        return " ".join(result_lines)

    # --- Fallback: CO-based format ---
    result_lines = []
    capturing = False

    for i, line in enumerate(lines):
        stripped = line.strip()
        low = stripped.lower()

        # Detect the table header row "CO  Topics to be discussed  Hrs."
        if not capturing and _CO_TABLE_HEADER_RE.match(stripped):
            capturing = True
            continue

        # Also start capturing on "CO1" if the NEXT line looks like content
        # (not a number, which would indicate the PO correlation table)
        if not capturing and _CO_SECTION_RE.match(stripped):
            # Peek ahead: if next non-empty line is text (not a digit), treat as content
            for j in range(i + 1, min(i + 4, len(lines))):
                nxt = lines[j].strip()
                if not nxt:
                    continue
                if nxt.isdigit() or re.match(r"^[A-Z]{1,3}\d*$", nxt):
                    break  # looks like PO table, skip
                # Looks like real content
                capturing = True
                break
            if capturing:
                continue  # skip the "CO1" label itself
            continue

        if capturing:
            if _line_starts_with_any(stripped, _UNIT_END_KEYWORDS):
                break
            # Skip CO labels (CO1, CO2, etc.)
            if _CO_SECTION_RE.match(stripped):
                continue
            # Skip hour counts (standalone numbers)
            if stripped.isdigit() or re.match(r"^\d+(\.\d+)?$", stripped):
                continue
            if not stripped:
                continue
            if low.startswith("total"):
                break
            if low.startswith("guest lecture"):
                break
            result_lines.append(stripped)

    return " ".join(result_lines)


def _extract_experiments(text: str) -> str:
    """Extract the list of experiments/lab section."""
    lines = text.split("\n")
    result_lines: list[str] = []
    capturing = False

    for line in lines:
        stripped = line.strip()

        if not capturing and _line_starts_with_any(stripped, _EXP_START_KEYWORDS):
            capturing = True
            continue

        if capturing:
            if _line_starts_with_any(stripped, _EXP_END_KEYWORDS):
                break
            # Skip pure numbers and empty lines
            if stripped.isdigit() or re.match(r"^\d+\.$", stripped) or not stripped:
                continue
            # Skip SO columns
            if re.match(r"^[a-l](?:\s*,\s*[a-l])*$", stripped):
                continue
            result_lines.append(stripped)

    return " ".join(result_lines)


def _extract_objectives(text: str) -> str:
    """Extract the course objectives section."""
    lines = text.split("\n")
    result_lines: list[str] = []
    capturing = False

    for line in lines:
        stripped = line.strip()

        if not capturing and _line_starts_with_any(stripped, _OBJ_START_KEYWORDS):
            capturing = True
            continue

        if capturing:
            if _line_starts_with_any(stripped, _OBJ_END_KEYWORDS):
                break
            if not stripped:
                continue
            result_lines.append(stripped)

    return " ".join(result_lines)


def _extract_course_header(text: str) -> Dict[str, Optional[str]]:
    """Extract course code and name from the top of the document."""
    lines = text.split("\n")
    code = None
    name_parts: list[str] = []
    code_re = re.compile(r"^[A-Z]{2,5}\d{3,4}[A-Z]?$")

    for i, line in enumerate(lines[:30]):  # only scan first 30 lines
        stripped = line.strip()
        if not stripped:
            continue

        # Skip the "Course Code" label
        if stripped.lower() in ("course code", "course name", "course title"):
            continue

        if code is None and code_re.match(stripped):
            code = stripped
            continue

        if code is not None and not name_parts:
            # Next meaningful non-metadata line is the course name
            if stripped.lower() in ("course type", "ct", "lt", "ltp", "lp", "pj"):
                break
            if re.match(r"^[A-Z]{1,3}\s+C$", stripped):  # "CT     C" pattern
                break
            if stripped.lower().startswith("credit"):
                break
            if re.match(r"^\d+(\.\d+)?$", stripped):
                break
            name_parts.append(stripped)

    return {
        "course_code": code,
        "course_name": " ".join(name_parts).strip() or None,
    }


# ──────────────────────────────────────────────────────────────────
# Public API
# ──────────────────────────────────────────────────────────────────

def extract_syllabus_from_text(raw_text: str) -> Dict:
    """
    Extract syllabus content from raw text of a course document.

    Returns
    -------
    dict with keys:
      - course_code (str | None)
      - course_name (str | None)
      - syllabus_text (str) — combined unit content + experiments
      - unit_content (str)
      - experiments (str)
      - objectives (str)
    """
    text = _clean_text(raw_text)
    header = _extract_course_header(text)
    unit_content = _extract_unit_content(text)
    experiments = _extract_experiments(text)
    objectives = _extract_objectives(text)

    # Build combined syllabus text
    parts = []
    if objectives:
        parts.append(f"Objectives: {objectives}")
    if unit_content:
        parts.append(f"Syllabus: {unit_content}")
    if experiments:
        parts.append(f"Experiments: {experiments}")

    syllabus_text = "\n".join(parts)

    return {
        "course_code": header["course_code"],
        "course_name": header["course_name"],
        "syllabus_text": syllabus_text,
        "unit_content": unit_content,
        "experiments": experiments,
        "objectives": objectives,
    }


def extract_syllabus_from_pdf(file_bytes: bytes) -> Dict:
    """
    Extract syllabus content from a course PDF file.

    Parameters
    ----------
    file_bytes : bytes
        Raw PDF content.

    Returns
    -------
    dict — see extract_syllabus_from_text.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF required: pip install PyMuPDF")

    full_text = ""
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            full_text += page.get_text("text") + "\n"

    return extract_syllabus_from_text(full_text)


def extract_syllabus_from_docx(file_bytes: bytes) -> Dict:
    """
    Extract syllabus content from a course DOCX file.

    Parameters
    ----------
    file_bytes : bytes
        Raw DOCX content.

    Returns
    -------
    dict — see extract_syllabus_from_text.
    """
    try:
        import docx
        import io
    except ImportError:
        raise RuntimeError("python-docx required: pip install python-docx")

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]

    # Also extract text from tables (common in course docs)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" ".join(cells))

    full_text = "\n".join(paragraphs)
    return extract_syllabus_from_text(full_text)


def extract_syllabus_from_file(filepath: str) -> Dict:
    """
    Dispatch to PDF or DOCX parser based on file extension.

    Parameters
    ----------
    filepath : str
        Path to the course PDF or DOCX file.

    Returns
    -------
    dict — see extract_syllabus_from_text.
    """
    with open(filepath, "rb") as f:
        data = f.read()

    if filepath.lower().endswith(".docx"):
        return extract_syllabus_from_docx(data)
    else:
        return extract_syllabus_from_pdf(data)
